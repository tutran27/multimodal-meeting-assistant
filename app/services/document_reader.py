"""
Module: document_reader.py
Vai trò: Service Layer chịu trách nhiệm xử lý và đọc nội dung văn bản từ các định dạng tệp tin khác nhau.

Mô tả chi tiết:
- Hỗ trợ phân tích cú pháp và trích xuất nội dung văn bản từ các định dạng phổ biến: văn bản thuần túy (.txt, .md), Microsoft Word (.docx) và PDF (.pdf).
- Sử dụng thư viện `python-docx` để đọc nội dung đoạn văn (paragraphs) cũng như thông tin dạng bảng biểu (tables) trong tài liệu Word.
- Sử dụng thư viện `pypdf` để đọc và trích xuất văn bản từ từng trang của tài liệu PDF.
- Hỗ trợ tự động chuẩn hóa và ghép nối nội dung từ nhiều trang khác nhau, loại bỏ các dòng trống thừa.
- Ném ra ngoại lệ `ValueError` đối với các định dạng tệp không được hỗ trợ để đảm bảo an toàn kiểu dữ liệu.
- Chứa hàm chạy thử nghiệm CLI (`__main__`) giúp kiểm tra nhanh kết quả đọc tệp tin bất kỳ thông qua đối số truyền vào từ dòng lệnh.
"""

from pathlib import Path
from docx import Document
from pypdf import PdfReader

class DocumentReader:
    def read(self, file_path: str | Path) -> str:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix == ".pdf":
            return self._read_pdf(path)
        raise ValueError(f"Unsupported script file: {suffix}")

    @staticmethod
    def _read_docx(path: Path) -> str:
        document = Document(path)
        lines = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(line for line in lines if line)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(page.strip() for page in pages if page.strip())

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python -m app.services.document_reader <file>")
    else:
        text = DocumentReader().read(sys.argv[1])
        print(text[:1000])